"""
File parsing utilities — extract text from PDF, DOCX, CSV, XLSX for the RAG pipeline.
"""

import logging
import os
from typing import Optional

import pandas as pd

logger = logging.getLogger(__name__)


def parse_file(file_path: str, file_type: str) -> str:
    """
    Dispatch to the correct parser based on file extension.
    Returns raw text suitable for chunking & embedding.
    """
    file_type = file_type.lower().strip(".")
    parsers = {
        "pdf": _parse_pdf,
        "docx": _parse_docx,
        "csv": _parse_csv,
        "xlsx": _parse_xlsx,
    }

    parser = parsers.get(file_type)
    if not parser:
        raise ValueError(f"Unsupported file type: {file_type}")

    text = parser(file_path)
    logger.info(
        "Parsed %s (%s): %d characters extracted",
        os.path.basename(file_path), file_type, len(text),
    )
    return text


# ── PDF ──────────────────────────────────────────────────────────────────────

def _parse_pdf(path: str) -> str:
    import fitz  # PyMuPDF

    doc = fitz.open(path)
    pages = []
    for i, page in enumerate(doc):
        text = page.get_text("text")
        if text.strip():
            pages.append(f"[Page {i + 1}]\n{text}")
    doc.close()
    return "\n\n".join(pages)


# ── DOCX ─────────────────────────────────────────────────────────────────────

def _parse_docx(path: str) -> str:
    from docx import Document

    doc = Document(path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

    # Also extract tables
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))

    return "\n\n".join(paragraphs)


# ── CSV ──────────────────────────────────────────────────────────────────────

def _parse_csv(path: str) -> str:
    try:
        df = pd.read_csv(path, encoding="utf-8")
    except UnicodeDecodeError:
        df = pd.read_csv(path, encoding="latin-1")

    return _dataframe_to_analysis(df, os.path.basename(path))


# ── XLSX ─────────────────────────────────────────────────────────────────────

def _parse_xlsx(path: str) -> str:
    xl = pd.ExcelFile(path)
    sections = []

    for sheet_name in xl.sheet_names:
        df = xl.parse(sheet_name)
        if df.empty:
            continue
        header = f"=== Sheet: {sheet_name} ==="
        analysis = _dataframe_to_analysis(df, sheet_name)
        sections.append(f"{header}\n{analysis}")

    return "\n\n".join(sections) if sections else "Empty spreadsheet"


# ── DataFrame → Rich Text ───────────────────────────────────────────────────

def _dataframe_to_analysis(df: pd.DataFrame, source_name: str) -> str:
    """
    Convert a DataFrame into a rich textual summary for LLM consumption.
    Includes statistics, correlations, distributions, and data quality notes.
    """
    parts = [f"Data Source: {source_name}", f"Shape: {df.shape[0]} rows × {df.shape[1]} columns"]

    # Column info
    parts.append(f"Columns: {', '.join(df.columns.tolist())}")

    # Data types
    dtype_summary = df.dtypes.value_counts().to_dict()
    parts.append(f"Data types: {dict(dtype_summary)}")

    # Missing values
    missing = df.isnull().sum()
    missing_cols = missing[missing > 0]
    if len(missing_cols) > 0:
        parts.append("Missing values:")
        for col, count in missing_cols.items():
            pct = count / len(df) * 100
            parts.append(f"  - {col}: {count} ({pct:.1f}%)")

    # Numeric statistics
    numeric_cols = df.select_dtypes(include=["number"]).columns.tolist()
    if numeric_cols:
        parts.append("\nNumeric Summary Statistics:")
        desc = df[numeric_cols].describe().round(2)
        parts.append(desc.to_string())

        # Correlations (if >1 numeric column)
        if len(numeric_cols) > 1:
            corr = df[numeric_cols].corr().round(3)
            parts.append("\nCorrelation Matrix:")
            parts.append(corr.to_string())

    # Categorical column summaries
    cat_cols = df.select_dtypes(include=["object", "category"]).columns.tolist()
    if cat_cols:
        parts.append("\nCategorical Columns:")
        for col in cat_cols[:10]:  # limit to first 10
            nunique = df[col].nunique()
            top = df[col].value_counts().head(5)
            parts.append(f"  {col} ({nunique} unique): {dict(top)}")

    # Sample rows (first 10)
    parts.append(f"\nSample Data (first {min(10, len(df))} rows):")
    parts.append(df.head(10).to_string(index=False))

    return "\n".join(parts)
